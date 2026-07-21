import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import networkx as nx
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
import matplotlib.image as mpimg
import numpy as np
import math
from matplotlib.patches import FancyArrow, Rectangle
import webbrowser
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from PIL import Image, ImageTk
import matplotlib.patheffects as pe

os.chdir(os.path.dirname(os.path.abspath(__file__)))

# --- Theme Configuration ---
current_mode = "light"

THEMES = {
    "light": {
        "bg_main": "#eef3f8",
        "bg_header": "#12355b",
        "fg_header_title": "#ffffff",
        "fg_header_sub": "#b9d7f5",
        "bg_card": "#ffffff",
        "border_card": "#d8e2ee",
        "fg_title": "#173b61",
        "fg_subtitle": "#627d98",
        "fg_label": "#334e68",
        "bg_input": "#ffffff",
        "fg_input": "#102a43",
        "bg_result": "#f8fafc",
        "fg_result": "#243b53",
        "btn_header_bg": "#e6f0fa",
        "btn_header_fg": "#12355b",
        "btn_header_active": "#d4e6f7",
        "plt_bg": "#ffffff",
        "plt_fg": "#173b61",
        "plt_edge": "#444444",
        "plt_title": "#173b61",
        "toggle_btn_text": "🌙 Dark Mode",
        "toggle_btn_bg": "#1e293b",
        "toggle_btn_fg": "#f8fafc",
    },
    "dark": {
        "bg_main": "#0f172a",
        "bg_header": "#1e293b",
        "fg_header_title": "#f8fafc",
        "fg_header_sub": "#94a3b8",
        "bg_card": "#1e293b",
        "border_card": "#334155",
        "fg_title": "#f1f5f9",
        "fg_subtitle": "#94a3b8",
        "fg_label": "#cbd5e1",
        "bg_input": "#334155",
        "fg_input": "#f8fafc",
        "bg_result": "#0f172a",
        "fg_result": "#e2e8f0",
        "btn_header_bg": "#334155",
        "btn_header_fg": "#f8fafc",
        "btn_header_active": "#475569",
        "plt_bg": "#1e293b",
        "plt_fg": "#f8fafc",
        "plt_edge": "#cbd5e1",
        "plt_title": "#f8fafc",
        "toggle_btn_text": "☀️ Light Mode",
        "toggle_btn_bg": "#f8fafc",
        "toggle_btn_fg": "#0f172a",
    }
}

# --- Tree layout (robust) ---
def hierarchy_pos(G, root, width=2., vert_gap=0.4, vert_loc=0, xcenter=0.5, pos=None, parent=None):
    if pos is None:
        pos = {root: (xcenter, vert_loc)}
    else:
        pos[root] = (xcenter, vert_loc)
    children = list(G.neighbors(root))
    if parent is not None and parent in children:
        children.remove(parent)
    if len(children) != 0:
        dx = width / len(children)
        nextx = xcenter - width / 2 - dx / 2
        for child in children:
            nextx += dx
            pos = hierarchy_pos(G, child, width=dx, vert_gap=vert_gap,
                                vert_loc=vert_loc - vert_gap, xcenter=nextx, pos=pos, parent=root)
    return pos


# Tooltip implementation
class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tipwindow = None
        widget.bind("<Enter>", self.show)
        widget.bind("<Leave>", self.hide)
    def show(self, event=None):
        if self.tipwindow or not self.text:
            return
        x = self.widget.winfo_rootx() + 25
        y = self.widget.winfo_rooty() + 20
        self.tipwindow = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(tw, text=self.text, justify=tk.LEFT,
                         background="#ffffe0", relief=tk.SOLID, borderwidth=1,
                         font=("tahoma", 9, "normal"))
        label.pack(ipadx=1)
    def hide(self, event=None):
        tw = self.tipwindow
        self.tipwindow = None
        if tw:
            tw.destroy()


def validate_int(P):
    if P == "": return True
    try: return int(P) > 0
    except ValueError: return False


def validate_float(P):
    if P == "": return True
    try: return float(P) >= 0
    except ValueError: return False


# Globals to save last run inputs/results
saved_graph = None
saved_topology = None
saved_ring = None
saved_result_text = ""
saved_positions = None


# Helper: draw double offset edges with opposite arrows (used for doubly-unidirectional ring)
def draw_double_edge_with_arrows(ax, x1, y1, x2, y2, color="#444"):
    dx = x2 - x1
    dy = y2 - y1
    L = math.hypot(dx, dy)
    if L == 0:
        return
    nxp = -dy / L
    nyp = dx / L
    offset = 0.06
    ax.plot([x1 + nxp*offset, x2 + nxp*offset], [y1 + nyp*offset, y2 + nyp*offset], color=color, linewidth=2, zorder=1)
    ax.plot([x1 - nxp*offset, x2 - nxp*offset], [y1 - nyp*offset, y2 - nyp*offset], color=color, linewidth=2, zorder=1)
    mid1x = (x1 + x2)/2 + nxp*offset
    mid1y = (y1 + y2)/2 + nyp*offset
    arr1 = FancyArrow(mid1x, mid1y, dx*0.0001, dy*0.0001, width=0.008, head_width=0.05, head_length=0.05, color=color, length_includes_head=True, zorder=3)
    ax.add_patch(arr1)
    mid2x = (x1 + x2)/2 - nxp*offset
    mid2y = (y1 + y2)/2 - nyp*offset
    arr2 = FancyArrow(mid2x, mid2y, -dx*0.0001, -dy*0.0001, width=0.008, head_width=0.05, head_length=0.05, color=color, length_includes_head=True, zorder=3)
    ax.add_patch(arr2)


def build_step_by_step(n, num_edges, port_cost, cable_len, cable_cost):
    lines = []
    lines.append("Step-by-step cost calculation:")
    lines.append(f"1) Number of nodes = {n}")
    lines.append(f"2) Number of connections (edges) = {num_edges}")
    lines.append(f"3) Cost per port = {port_cost}")
    lines.append(f"4) Total port cost = connections * 2 * cost_per_port")
    total_port_cost = num_edges * 2 * port_cost
    lines.append(f"   = {num_edges} * 2 * {port_cost} = {total_port_cost:.2f}")
    lines.append(f"5) Cable length per connection = {cable_len} m")
    lines.append(f"6) Cost per unit cable = {cable_cost} per m")
    total_cable_cost = num_edges * cable_len * cable_cost
    lines.append(f"7) Total cable cost = connections * cable_len * cable_cost")
    lines.append(f"   = {num_edges} * {cable_len} * {cable_cost} = {total_cable_cost:.2f}")
    total_cost = total_port_cost + total_cable_cost
    lines.append(f"8) Total cost = total_port_cost + total_cable_cost = {total_port_cost:.2f} + {total_cable_cost:.2f} = {total_cost:.2f}")
    return "\n".join(lines), total_port_cost, total_cable_cost, total_cost


def generate_topology():
    global saved_graph, saved_topology, saved_ring, saved_result_text, saved_positions
    try:
        n = int(entry_nodes.get())
    except ValueError:
        messagebox.showerror("Input Error", "Enter a positive integer for number of nodes.")
        return

    if n < 2:
        messagebox.showerror("Input Error", "Use at least 2 nodes to create a network topology.")
        return
    if n > 50:
        messagebox.showerror("Input Error", "Use 50 nodes or fewer so the diagram stays readable.")
        return

    topology_type = topology_choice.get()
    ring_variant_val = ring_variant.get() if topology_type == "Ring" else None

    if topology_type == "Ring" and ring_variant_val in ["Singly (Unidirectional)", "Doubly (Unidirectional)"]:
        G = nx.DiGraph()
    else:
        G = nx.Graph()
    G.add_nodes_from(range(1, n + 1))

    if topology_type == "Bus":
        for i in range(1, n):
            G.add_edge(i, i + 1)
    elif topology_type == "Star":
        for i in range(2, n + 1):
            G.add_edge(1, i)
    elif topology_type == "Ring":
        for i in range(1, n):
            G.add_edge(i, i + 1)
        G.add_edge(n, 1)
        if ring_variant_val == "Doubly (Unidirectional)":
            for i in range(1, n):
                G.add_edge(i + 1, i)
            G.add_edge(1, n)
    elif topology_type == "Mesh":
        for i in range(1, n + 1):
            for j in range(i + 1, n + 1):
                G.add_edge(i, j)
    elif topology_type == "Tree":
        k = math.log2(n + 1)
        if not k.is_integer():
            messagebox.showerror("Tree Requirement", "Number of nodes for Tree must be 2^k - 1 (e.g., 3,7,15...) for a perfect binary tree.")
            return
        for i in range(2, n + 1):
            G.add_edge(i // 2, i)

    try:
        port_cost = float(entry_port_cost.get())
        cable_len = float(entry_cable_length.get())
        cable_cost = float(entry_cable_cost.get())
    except ValueError:
        messagebox.showerror("Input Error", "Please check cost values and lengths.")
        return

    if not all(math.isfinite(value) and value >= 0 for value in (port_cost, cable_len, cable_cost)):
        messagebox.showerror("Input Error", "Cost values and cable length must be finite, non-negative numbers.")
        return

    num_edges = G.number_of_edges()
    step_text, total_port_cost, total_cable_cost, total_cost = build_step_by_step(n, num_edges, port_cost, cable_len, cable_cost)

    summary = (
        f"Topology: {topology_type} {f'({ring_variant_val})' if ring_variant_val else ''}\n"
        f"Nodes: {n}\n"
        f"Connections: {num_edges}\n"
        f"Total Cost: ₹{total_cost:.2f}\n\n"
        f"{step_text}\n"
    )

    result_label.config(state='normal')
    result_label.delete("1.0", tk.END)
    result_label.insert(tk.END, summary)
    result_label.config(state='disabled')

    saved_graph = G
    saved_topology = topology_type
    saved_ring = ring_variant_val
    saved_result_text = summary
    download_button.config(state="normal")

    show_graph(G, topology_type, ring_variant_val)


def show_graph(G, topology_type, ring_variant_val, preview=False):
    global saved_positions
    t = THEMES[current_mode]
    for widget in graph_frame.winfo_children():
        widget.destroy()
    fig, ax = plt.subplots(figsize=(9, 7))
    fig.patch.set_facecolor(t["plt_bg"])
    ax.set_facecolor(t["plt_bg"])

    # --- Layout setup ---
    if topology_type == "Bus":
        nodes = list(G.nodes)
        n = len(nodes)
        xs = np.arange(n)
        ys = np.array([0.6 if i % 2 == 0 else -0.6 for i in range(n)])
        pos = {node: (float(x), float(y)) for node, x, y in zip(nodes, xs, ys)}
        left, right = -0.5, n - 0.5
        ax.plot([left, right], [0, 0], color=t["plt_edge"], linewidth=4, zorder=0)
        # Terminators at both ends
        for tx in [left, right]:
            term_rect = Rectangle((tx - 0.09, -0.11), 0.18, 0.22, linewidth=0, edgecolor=None, facecolor=t["plt_edge"], zorder=2)
            ax.add_patch(term_rect)
        for x, y in zip(xs, ys):
            ax.plot([x, x], [y - 0.08 if y > 0 else y + 0.08, 0], color=t["plt_edge"], linewidth=2, zorder=1)
            ax.plot([x - 0.06, x + 0.06], [0, 0], color=t["plt_edge"], linewidth=2, zorder=1)
    elif topology_type == "Ring":
        n = len(G.nodes)
        theta = np.linspace(0, 2 * np.pi, n, endpoint=False)
        pos = {node: (np.cos(th), np.sin(th)) for node, th in zip(G.nodes, theta)}
    elif topology_type == "Tree":
        try:
            pos = hierarchy_pos(G, 1)
        except Exception:
            pos = nx.spring_layout(G, seed=42)
    else:
        pos = nx.spring_layout(G, seed=42)

    saved_positions = pos.copy()

    # --- Draw Edges ---
    if topology_type != "Bus":
        edges_drawn = set()
        for u, v in G.edges():
            x1, y1 = pos[u]
            x2, y2 = pos[v]
            if topology_type == "Ring" and ring_variant_val == "Doubly (Unidirectional)":
                pair = tuple(sorted((u, v)))
                if pair not in edges_drawn:
                    draw_double_edge_with_arrows(ax, x1, y1, x2, y2, color=t["plt_edge"])
                    edges_drawn.add(pair)
            else:
                ax.plot([x1, x2], [y1, y2], color=t["plt_edge"], linewidth=2, zorder=1)
                if topology_type == "Ring" and ring_variant_val == "Singly (Unidirectional)":
                    dx = x2 - x1
                    dy = y2 - y1
                    midpoint = (x1 + 0.5 * dx, y1 + 0.5 * dy)
                    arr = FancyArrow(midpoint[0], midpoint[1], dx * 0.0001, dy * 0.0001,
                                     width=0.008, head_width=0.05, head_length=0.05,
                                     color=t["plt_edge"], length_includes_head=True, zorder=3)
                    ax.add_patch(arr)

    # --- Draw Nodes as Computer Icons ---
    try:
        img = mpimg.imread("computer.png")
        node_count = len(pos)
        zoom = 0.12 if node_count <= 8 else (0.09 if node_count <= 16 else 0.06)
        for node, (x, y) in pos.items():
            ab = AnnotationBbox(OffsetImage(img, zoom=zoom), (x, y), frameon=False, zorder=4)
            ax.add_artist(ab)
            ax.text(x, y + 0.02, str(node),
                    fontsize=12, fontweight='bold',
                    ha='center', va='center',
                    color='white',
                    path_effects=[pe.withStroke(linewidth=2, foreground='black')],
                    zorder=6)
    except Exception:
        nx.draw_networkx_nodes(G, pos, ax=ax, node_size=1100, node_color="#b7e2f9")
        nx.draw_networkx_labels(G, pos, ax=ax, font_size=13, font_weight='bold')

    ax.set_axis_off()
    plt.title(f"{topology_type}" + (f" ({ring_variant_val})" if ring_variant_val else ""), fontsize=15, weight='bold', color=t["plt_title"])

    xs = [p[0] for p in pos.values()]
    ys = [p[1] for p in pos.values()]
    if xs and ys:
        xmin, xmax = min(xs) - 0.9, max(xs) + 0.9
        ymin, ymax = min(ys) - 0.9, max(ys) + 0.9
        ax.set_xlim(xmin, xmax)
        ax.set_ylim(ymin, ymax)

    canvas = FigureCanvasTkAgg(fig, master=graph_frame)
    canvas.draw()
    canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
    toolbar = NavigationToolbar2Tk(canvas, graph_frame)
    toolbar.update()
    toolbar.pack(side=tk.BOTTOM, fill=tk.X)
    try:
        toolbar.config(background=t["bg_card"])
        for child in toolbar.winfo_children():
            child.config(background=t["bg_card"])
    except Exception:
        pass
    plt.close(fig)


# Developed By: Full-size dialog with theme support and image fallback
def show_developed_by():
    t = THEMES[current_mode]
    win = tk.Toplevel(root)
    win.title("Developed By - Project Credits")
    win.geometry("640x540")
    win.minsize(520, 440)
    win.resizable(True, True)
    win.configure(bg=t["bg_main"])
    win.transient(root)

    # Header bar
    hdr = tk.Frame(win, bg=t["bg_header"], padx=28, pady=20)
    hdr.pack(fill=tk.X)
    tk.Label(hdr, text="Project Credits", font=("Segoe UI", 18, "bold"),
             fg=t["fg_header_title"], bg=t["bg_header"]).pack(anchor="w")
    tk.Label(hdr, text="Network Topology Simulator & Cost Estimator", font=("Segoe UI", 10),
             fg=t["fg_header_sub"], bg=t["bg_header"]).pack(anchor="w", pady=(3, 0))

    container = tk.Frame(win, bg=t["bg_main"], padx=24, pady=20)
    container.pack(fill=tk.BOTH, expand=True)

    contributors = [
        ("Abijith Thennarasu", "24BCE1626", "Lead Developer", "abijith.jpeg", "👨‍💻"),
        ("Dharmayu Jadwani", "Project Contributor", "Developer", "dharmyu.jpeg", "👨‍💻"),
        ("Dr. Swaminathan Annadurai", "Faculty Guide", "Project Advisor", "professor.png", "👨‍🏫"),
    ]

    for name, detail, role, img_name, icon in contributors:
        card = tk.Frame(container, bg=t["bg_card"], highlightbackground=t["border_card"], highlightthickness=1, padx=18, pady=14)
        card.pack(fill=tk.X, pady=8)

        avatar_frame = tk.Frame(card, bg=t["bg_card"])
        avatar_frame.pack(side=tk.LEFT, padx=(0, 18))

        img_obj = None
        if os.path.exists(img_name):
            try:
                pil_img = Image.open(img_name).resize((54, 54), Image.Resampling.LANCZOS)
                img_obj = ImageTk.PhotoImage(pil_img)
            except Exception:
                img_obj = None

        if img_obj:
            lbl_img = tk.Label(avatar_frame, image=img_obj, bg=t["bg_card"])
            lbl_img.image = img_obj
            lbl_img.pack()
        else:
            tk.Label(avatar_frame, text=icon, font=("Segoe UI Emoji", 26), bg=t["bg_card"]).pack()

        info_frame = tk.Frame(card, bg=t["bg_card"])
        info_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        tk.Label(info_frame, text=name, font=("Segoe UI", 13, "bold"), fg=t["fg_title"], bg=t["bg_card"]).pack(anchor="w")
        tk.Label(info_frame, text=f"{role}  |  {detail}", font=("Segoe UI", 9.5), fg=t["fg_subtitle"], bg=t["bg_card"]).pack(anchor="w", pady=(3, 0))

    btm = tk.Frame(win, bg=t["bg_main"], pady=14)
    btm.pack(fill=tk.X)
    tk.Button(btm, text="Close", command=win.destroy, font=("Segoe UI", 10, "bold"),
              bg="#1d4ed8", fg="white", activebackground="#1e40af", activeforeground="white",
              relief="flat", padx=26, pady=8, cursor="hand2").pack()


def show_learn():
    t = THEMES[current_mode]
    win = tk.Toplevel(root)
    win.title("Learn - Network Topology Concepts")
    win.geometry("950x700")
    win.config(bg=t["bg_main"])

    tk.Label(win, text="Network Topology Concepts", font=("Segoe UI", 16, "bold"),
             bg=t["bg_main"], fg="#007ACC").pack(pady=(12, 8))

    frame = tk.Frame(win, bg=t["bg_main"])
    frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

    scrollbar = tk.Scrollbar(frame)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    text_widget = tk.Text(frame, wrap=tk.WORD, font=("Segoe UI", 10),
                          bg=t["bg_card"], fg=t["fg_input"], yscrollcommand=scrollbar.set,
                          padx=15, pady=10)
    text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    scrollbar.config(command=text_widget.yview)

    learn_content = """
NETWORK TOPOLOGY

═══════════════════════════════════════════════════════════════════════

1. INTRODUCTION TO NETWORK TOPOLOGY

Network topology refers to the arrangement of different elements (links, nodes, etc.) in a computer network. It defines how devices are connected and how data flows between them.

Types of topology:
• Physical Topology: Actual layout of cables and devices
• Logical Topology: How data actually flows in the network

═══════════════════════════════════════════════════════════════════════

2. BUS TOPOLOGY

Structure:
• All devices connect to a single central cable (backbone)
• Data travels in both directions along the cable
• Terminators at both ends prevent signal reflection

Advantages:
✓ Easy to install and extend
✓ Cost-effective (requires less cable)
✓ Suitable for small networks
✓ Easy to understand and implement

Disadvantages:
✗ Single point of failure (if backbone fails, entire network fails)
✗ Performance degrades with more devices
✗ Difficult to troubleshoot
✗ Limited cable length and number of stations

Best Use Cases:
• Small networks (10-12 computers)
• Temporary networks
• Budget-constrained setups

═══════════════════════════════════════════════════════════════════════

3. STAR TOPOLOGY

Structure:
• All devices connect to a central hub/switch
• Each device has dedicated point-to-point connection to hub
• Data passes through the central hub

Advantages:
✓ Easy to install and manage
✓ Failure of one node doesn't affect others
✓ Easy to detect faults and troubleshoot
✓ Easy to add/remove devices
✓ Better performance than bus topology

Disadvantages:
✗ Hub/switch is single point of failure
✗ Requires more cable than bus topology
✗ Cost increases with hub/switch quality
✗ Limited by hub specifications

Best Use Cases:
• Small to medium office networks
• Home networks
• Networks requiring easy management

═══════════════════════════════════════════════════════════════════════

4. RING TOPOLOGY

Structure:
• Each device connects to exactly two other devices
• Forms a circular data path
• Data travels in one or both directions

Variants:
a) Singly Bidirectional: Data can flow both ways on single ring
b) Singly Unidirectional: Data flows in one direction only
c) Doubly Unidirectional: Two separate rings with opposite data flows

Advantages:
✓ Equal access for all devices
✓ No collisions (token passing)
✓ Predictable performance
✓ Can span longer distances than bus

Disadvantages:
✗ Failure of one device/cable affects entire network
✗ Difficult to troubleshoot
✗ Adding/removing devices disrupts network
✗ More expensive than bus topology

Best Use Cases:
• Networks requiring equal access
• Token ring networks (legacy)
• High-speed LANs with token passing

═══════════════════════════════════════════════════════════════════════

5. MESH TOPOLOGY

Structure:
• Every device connects to every other device
• Multiple paths between any two devices
• Provides redundancy and fault tolerance

Types:
a) Full Mesh: Every device connected to every other device
   Connections = n(n-1)/2 where n = number of nodes
   
b) Partial Mesh: Some devices have multiple connections, not all

Advantages:
✓ Highly reliable (multiple paths)
✓ No traffic congestion
✓ Robust and fault-tolerant
✓ Data can be transmitted simultaneously
✓ Privacy and security

Disadvantages:
✗ Very expensive (many cables and ports needed)
✗ Complex installation and maintenance
✗ Requires large amount of cabling
✗ Difficult to configure

Best Use Cases:
• Critical systems requiring high reliability
• Military and financial networks
• Backbone networks
• Wireless mesh networks

═══════════════════════════════════════════════════════════════════════

6. TREE TOPOLOGY

Structure:
• Hierarchical structure like a tree
• Combination of star topologies
• Root node connects to secondary nodes, which connect to tertiary nodes
• Also called Hierarchical Topology

Advantages:
✓ Scalable (easy to expand)
✓ Easy to manage and maintain
✓ Error detection is easy
✓ Suitable for large networks
✓ Point-to-point wiring for segments

Disadvantages:
✗ If backbone fails, entire segment fails
✗ More cables required
✗ Expensive
✗ Difficult to configure

Best Use Cases:
• Large organizations
• Wide Area Networks (WANs)
• Networks with hierarchical structure
• Educational institutions

═══════════════════════════════════════════════════════════════════════

7. COST ANALYSIS FACTORS

When selecting a topology, consider these costs:

1. Port/Interface Costs:
   • Network interface cards (NICs)
   • Switch/hub ports
   • Each connection requires 2 ports

2. Cable Costs:
   • Cable length × cost per unit
   • Different cable types (CAT5e, CAT6, fiber)
   • Installation labor

3. Hardware Costs:
   • Central devices (switches, hubs, routers)
   • Quality and features affect price
   • Redundant devices for fault tolerance

4. Maintenance Costs:
   • Troubleshooting complexity
   • Downtime costs
   • Replacement parts

═══════════════════════════════════════════════════════════════════════

8. COMPARISON SUMMARY

Topology    | Cost    | Reliability | Scalability | Complexity
------------|---------|-------------|-------------|------------
Bus         | Low     | Low         | Poor        | Simple
Star        | Medium  | Medium      | Good        | Simple
Ring        | Medium  | Medium      | Medium      | Medium
Mesh        | High    | Very High   | Poor        | Complex
Tree        | High    | Medium      | Excellent   | Medium

═══════════════════════════════════════════════════════════════════════

9. SELECTION CRITERIA

Choose topology based on:

✓ Network size (number of devices)
✓ Budget constraints
✓ Reliability requirements
✓ Scalability needs
✓ Performance requirements
✓ Ease of maintenance
✓ Physical constraints
✓ Security requirements

═══════════════════════════════════════════════════════════════════════

10. PRACTICAL CONSIDERATIONS

Installation Tips:
• Plan cable routes before installation
• Leave room for expansion
• Label all cables and ports
• Document network layout
• Test connections before deployment
• Consider environmental factors
• Plan for redundancy in critical systems

Troubleshooting:
• Bus: Check terminators and backbone
• Star: Check central hub/switch first
• Ring: Locate break in the ring
• Mesh: Test individual connections
• Tree: Check hierarchical levels systematically

═══════════════════════════════════════════════════════════════════════

REFERENCE VIDEO:
Animated Network Topology Explanation
https://www.youtube.com/watch?v=zbqrNg4C98U

For detailed examples and calculations, use the simulator to experiment with different topologies and cost parameters.

═══════════════════════════════════════════════════════════════════════
    """
    text_widget.insert('1.0', learn_content)
    text_widget.config(state='disabled')

    tk.Label(win, text="Reference Video (YouTube):", 
             font=("Segoe UI", 12, "bold"), bg=t["bg_main"], fg="#007ACC").pack(pady=(8, 6))

    def open_youtube_video():
        webbrowser.open("https://www.youtube.com/watch?v=zbqrNg4C98U")

    youtube_btn = tk.Button(
        win, text="▶  Watch on YouTube", 
        command=open_youtube_video,
        font=("Segoe UI", 11, "bold"), 
        bg="#FF0000", fg="white",
        padx=15, pady=8,
        relief="flat", cursor="hand2"
    )
    youtube_btn.pack(pady=(5, 15))

    tk.Button(win, text="Close", command=win.destroy,
              font=("Segoe UI", 10), padx=20, pady=5).pack(pady=10)


def save_graph_to_file(filepath):
    if saved_graph is None or saved_positions is None:
        return False
    
    G = saved_graph
    topology_type = saved_topology
    ring_variant_val = saved_ring
    pos = saved_positions
    t = THEMES[current_mode]
    
    fig = plt.figure(figsize=(10, 8))
    ax = fig.add_subplot(111)
    fig.patch.set_facecolor(t["plt_bg"])
    ax.set_facecolor(t["plt_bg"])
    
    if topology_type == "Bus":
        left, right = -0.5, len(pos) - 0.5
        ax.plot([left, right], [0, 0], color=t["plt_edge"], linewidth=4, zorder=0)
        for tx in [left, right]:
            term_rect = Rectangle((tx - 0.09, -0.11), 0.18, 0.22, linewidth=0, edgecolor=None, facecolor=t["plt_edge"], zorder=2)
            ax.add_patch(term_rect)
        for node, (x, y) in pos.items():
            ax.plot([x, x], [y - 0.08 if y > 0 else y + 0.08, 0], color=t["plt_edge"], linewidth=2)
            ax.plot([x - 0.06, x + 0.06], [0, 0], color=t["plt_edge"], linewidth=2)
    
    edges_drawn = set()
    if topology_type != "Bus":
        for u, v in G.edges():
            x1, y1 = pos[u]; x2, y2 = pos[v]
            if topology_type == "Ring" and ring_variant_val == "Doubly (Unidirectional)":
                pair = tuple(sorted((u, v)))
                if pair not in edges_drawn:
                    draw_double_edge_with_arrows(ax, x1, y1, x2, y2, color=t["plt_edge"])
                    edges_drawn.add(pair)
            else:
                ax.plot([x1, x2], [y1, y2], color=t["plt_edge"], linewidth=2)
                if topology_type == "Ring" and ring_variant_val == "Singly (Unidirectional)":
                    dx = x2 - x1; dy = y2 - y1
                    mid = (x1 + 0.5*dx, y1 + 0.5*dy)
                    arr = FancyArrow(mid[0], mid[1], dx*0.0001, dy*0.0001, width=0.008, head_width=0.05, head_length=0.05, color=t["plt_edge"], length_includes_head=True)
                    ax.add_patch(arr)
    
    try:
        img = mpimg.imread("computer.png")
        node_count = len(pos)
        zoom = 0.14 if node_count <= 8 else (0.1 if node_count <= 16 else 0.07)
        for node, (x, y) in pos.items():
            ab = AnnotationBbox(OffsetImage(img, zoom=zoom), (x, y), frameon=False)
            ax.add_artist(ab)
    except Exception:
        nx.draw_networkx_nodes(G, pos, ax=ax, node_size=1100, node_color="#b7e2f9")
    
    if topology_type == "Bus":
        for node, (x, y) in pos.items():
            ypos = y + 0.15 if y > 0 else y - 0.15
            ax.text(x, ypos, str(node), fontsize=12, fontweight='bold', ha='center', va='center')
    else:
        nx.draw_networkx_labels(G, pos, ax=ax, font_size=13, font_weight='bold')
    
    ax.set_axis_off()
    plt.title(f"{topology_type}" + (f" ({ring_variant_val})" if ring_variant_val else ""), fontsize=15, weight='bold', color=t["plt_title"])
    
    fig.savefig(filepath, bbox_inches='tight', dpi=150)
    plt.close(fig)
    return True


def do_download():
    if saved_graph is None:
        messagebox.showinfo("Nothing to download", "Generate a topology first, then download.")
        return
    
    try:
        fpath = filedialog.asksaveasfilename(
            defaultextension=".docx", 
            filetypes=[("Word Document","*.docx")], 
            title="Save detailed report as..."
        )
        if not fpath:
            return
        
        doc = Document()
        title = doc.add_heading('Network Topology Simulator - Detailed Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format = title.runs[0].font
        title_format.color.rgb = RGBColor(0, 122, 204)
        
        doc.add_paragraph()
        doc.add_heading('Input Parameters', 1)
        topology_type = saved_topology or ""
        ring_variant_val = saved_ring or ""
        
        input_table = doc.add_table(rows=6, cols=2)
        input_table.style = 'Light Grid Accent 1'
        
        input_data = [
            ('Topology Type', f"{topology_type} {('('+ring_variant_val+')') if ring_variant_val else ''}"),
            ('Number of Nodes', entry_nodes.get()),
            ('Cost per Port (₹)', entry_port_cost.get()),
            ('Cable Length per Connection (m)', entry_cable_length.get()),
            ('Cost per Unit Cable (₹/m)', entry_cable_cost.get()),
            ('Total Connections', str(saved_graph.number_of_edges()))
        ]
        
        for i, (param, value) in enumerate(input_data):
            input_table.rows[i].cells[0].text = param
            input_table.rows[i].cells[1].text = value
            input_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        doc.add_heading('Cost Analysis Results', 1)
        
        result_para = doc.add_paragraph()
        result_para.add_run(saved_result_text).font.name = 'Consolas'
        
        doc.add_paragraph()
        doc.add_heading('Network Topology Diagram', 1)
        
        import tempfile
        temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        temp_img_path = temp_img.name
        temp_img.close()
        
        if save_graph_to_file(temp_img_path):
            try:
                doc.add_picture(temp_img_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Error embedding graph image: {e}]")
            finally:
                try:
                    os.unlink(temp_img_path)
                except:
                    pass
        else:
            doc.add_paragraph("[Graph image could not be generated]")
        
        doc.add_paragraph()
        doc.add_heading('Network Connections', 1)
        edges_para = doc.add_paragraph()
        
        for u, v in saved_graph.edges():
            edges_para.add_run(f"• Node {u} ↔ Node {v}\n")
        
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run('\n' + '─' * 70 + '\n').font.color.rgb = RGBColor(128, 128, 128)
        footer_run = footer_para.add_run('Generated by Network Topology Simulator\n')
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(fpath)
        messagebox.showinfo("Success", f"Detailed report with graph saved to:\n{fpath}")
        
    except Exception as e:
        messagebox.showerror("Error", f"Could not create report: {e}\n\nMake sure python-docx is installed:\npip install python-docx")


# ------------------- GUI Initialization -------------------
root = tk.Tk()
root.title("Network Topology Simulator & Cost Estimator")
root.geometry("1240x800")
root.minsize(980, 680)

style = ttk.Style(root)
style.theme_use("clam")
style.configure("TCombobox", padding=6, font=("Segoe UI", 10))

# Application Header
header_bar = tk.Frame(root, height=72)
header_bar.pack(fill=tk.X)
header_bar.pack_propagate(False)

lbl_hdr_title = tk.Label(header_bar, text="Network Topology Simulator", font=("Segoe UI", 20, "bold"))
lbl_hdr_title.pack(side=tk.LEFT, padx=28, pady=(13, 0))

lbl_hdr_sub = tk.Label(header_bar, text="Design, visualize, and estimate your network", font=("Segoe UI", 10))
lbl_hdr_sub.pack(side=tk.LEFT, padx=(0, 20), pady=(20, 0))

btn_frame = tk.Frame(header_bar)
btn_frame.pack(side=tk.RIGHT, padx=22, pady=16)

btn_learn = tk.Button(btn_frame, text="Learn", command=show_learn, font=("Segoe UI", 9, "bold"),
                      relief="flat", padx=14, pady=6, cursor="hand2")
btn_learn.pack(side=tk.LEFT, padx=4)

btn_credits = tk.Button(btn_frame, text="Credits", command=show_developed_by, font=("Segoe UI", 9, "bold"),
                        relief="flat", padx=14, pady=6, cursor="hand2")
btn_credits.pack(side=tk.LEFT, padx=4)

btn_theme = tk.Button(btn_frame, text="🌙 Dark Mode", command=lambda: toggle_theme(), font=("Segoe UI", 9, "bold"),
                      relief="flat", padx=14, pady=6, cursor="hand2")
btn_theme.pack(side=tk.LEFT, padx=4)

download_button = tk.Button(btn_frame, text="Download report", command=do_download,
                            font=("Segoe UI", 9, "bold"), bg="#20a36a", fg="white",
                            activebackground="#188556", activeforeground="white", relief="flat", padx=14, pady=6,
                            cursor="hand2", state="disabled")
download_button.pack(side=tk.LEFT, padx=(8, 0))

main_frame = tk.Frame(root)
main_frame.pack(fill=tk.BOTH, expand=True, padx=18, pady=18)

# Left Control Panel
input_panel = tk.Frame(main_frame, width=390, highlightthickness=1)
input_panel.pack(side=tk.LEFT, fill=tk.Y)
input_panel.pack_propagate(False)

input_title_lbl = tk.Label(input_panel, text="Build your network", font=("Segoe UI", 17, "bold"))
input_title_lbl.pack(anchor="w", padx=22, pady=(22, 2))

input_sub_lbl = tk.Label(input_panel, text="Choose a topology and cost assumptions.", font=("Segoe UI", 10))
input_sub_lbl.pack(anchor="w", padx=22, pady=(0, 16))

frame_main = ttk.LabelFrame(input_panel, text="  Network design  ", style="Card.TLabelframe")
frame_main.pack(fill=tk.X, padx=18, pady=(0, 12))
frame_main.columnconfigure(1, weight=1)

vcmd_int = root.register(validate_int)
vcmd_float = root.register(validate_float)

form_labels = []
for row, label in enumerate(("Number of nodes", "Topology", "Ring variant")):
    lbl = tk.Label(frame_main, text=label, font=("Segoe UI", 10))
    lbl.grid(row=row, column=0, sticky="w", padx=12, pady=8)
    form_labels.append(lbl)

entry_nodes = tk.Entry(frame_main, width=12, validate="key", validatecommand=(vcmd_int, "%P"),
                       font=("Segoe UI", 10), relief="solid", bd=1)
entry_nodes.grid(row=0, column=1, sticky="ew", padx=(6, 12), pady=8)
entry_nodes.insert(0, "6")
ToolTip(entry_nodes, "Enter 2–50 nodes. Tree diagrams require 3, 7, 15, or 31 nodes.")

topology_choice = ttk.Combobox(frame_main, values=["Bus", "Star", "Ring", "Mesh", "Tree"],
                                font=("Segoe UI", 10), state="readonly")
topology_choice.grid(row=1, column=1, sticky="ew", padx=(6, 12), pady=8)
topology_choice.current(0)

ring_variant = ttk.Combobox(frame_main, values=[
    "Singly (Bidirectional)", "Singly (Unidirectional)", "Doubly (Unidirectional)"
], font=("Segoe UI", 10), state="disabled")
ring_variant.grid(row=2, column=1, sticky="ew", padx=(6, 12), pady=8)
ring_variant.current(0)

def update_ring_dropdown(event=None):
    ring_variant.config(state="readonly" if topology_choice.get() == "Ring" else "disabled")

topology_choice.bind("<<ComboboxSelected>>", update_ring_dropdown)

frame_cost = ttk.LabelFrame(input_panel, text="  Cost assumptions  ", style="Card.TLabelframe")
frame_cost.pack(fill=tk.X, padx=18, pady=(0, 12))
frame_cost.columnconfigure(1, weight=1)

for row, label in enumerate(("Cost per port (Rs.)", "Cable per link (m)", "Cable cost (Rs./m)")):
    lbl = tk.Label(frame_cost, text=label, font=("Segoe UI", 10))
    lbl.grid(row=row, column=0, sticky="w", padx=12, pady=8)
    form_labels.append(lbl)

entry_port_cost = tk.Entry(frame_cost, validate="key", validatecommand=(vcmd_float, "%P"), font=("Segoe UI", 10), relief="solid", bd=1)
entry_port_cost.grid(row=0, column=1, sticky="ew", padx=(6, 12), pady=8)
entry_port_cost.insert(0, "100")

entry_cable_length = tk.Entry(frame_cost, validate="key", validatecommand=(vcmd_float, "%P"), font=("Segoe UI", 10), relief="solid", bd=1)
entry_cable_length.grid(row=1, column=1, sticky="ew", padx=(6, 12), pady=8)
entry_cable_length.insert(0, "10")

entry_cable_cost = tk.Entry(frame_cost, validate="key", validatecommand=(vcmd_float, "%P"), font=("Segoe UI", 10), relief="solid", bd=1)
entry_cable_cost.grid(row=2, column=1, sticky="ew", padx=(6, 12), pady=8)
entry_cable_cost.insert(0, "50")

btn_generate = tk.Button(input_panel, text="Generate topology", command=generate_topology, font=("Segoe UI", 11, "bold"),
                         bg="#1d4ed8", fg="white", activebackground="#1e40af", activeforeground="white", relief="flat",
                         padx=20, pady=10, cursor="hand2")
btn_generate.pack(fill=tk.X, padx=18, pady=(0, 12))

frame_results = ttk.LabelFrame(input_panel, text="  Cost summary  ", style="Card.TLabelframe")
frame_results.pack(fill=tk.BOTH, expand=True, padx=18, pady=(0, 18))
result_label = tk.Text(frame_results, font=("Consolas", 9), height=12,
                       state="disabled", bd=0, wrap="word", padx=12, pady=10, cursor="arrow")
result_label.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)

# Right Visualization Panel
graph_card = tk.Frame(main_frame, highlightthickness=1)
graph_card.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(16, 0))

graph_card_title = tk.Label(graph_card, text="Topology visualization", font=("Segoe UI", 16, "bold"))
graph_card_title.pack(anchor="w", padx=22, pady=(20, 0))

graph_card_sub = tk.Label(graph_card, text="Generate a design to see its connection map.", font=("Segoe UI", 10))
graph_card_sub.pack(anchor="w", padx=22, pady=(2, 12))

graph_frame = tk.Frame(graph_card)
graph_frame.pack(fill=tk.BOTH, expand=True, padx=12, pady=(0, 12))

graph_placeholder = tk.Label(graph_frame, text="Your network diagram will appear here", font=("Segoe UI", 13))
graph_placeholder.place(relx=0.5, rely=0.45, anchor="center")

# --- Theme Application Logic ---
def apply_theme():
    t = THEMES[current_mode]
    root.configure(bg=t["bg_main"])
    header_bar.configure(bg=t["bg_header"])
    lbl_hdr_title.configure(bg=t["bg_header"], fg=t["fg_header_title"])
    lbl_hdr_sub.configure(bg=t["bg_header"], fg=t["fg_header_sub"])
    btn_frame.configure(bg=t["bg_header"])

    btn_learn.configure(bg=t["btn_header_bg"], fg=t["btn_header_fg"], activebackground=t["btn_header_active"])
    btn_credits.configure(bg=t["btn_header_bg"], fg=t["btn_header_fg"], activebackground=t["btn_header_active"])
    btn_theme.configure(text=t["toggle_btn_text"], bg=t["toggle_btn_bg"], fg=t["toggle_btn_fg"],
                        activebackground=t["toggle_btn_bg"], activeforeground=t["toggle_btn_fg"])

    main_frame.configure(bg=t["bg_main"])
    input_panel.configure(bg=t["bg_card"], highlightbackground=t["border_card"])
    input_title_lbl.configure(bg=t["bg_card"], fg=t["fg_title"])
    input_sub_lbl.configure(bg=t["bg_card"], fg=t["fg_subtitle"])

    style.configure("Card.TLabelframe", background=t["bg_card"], bordercolor=t["border_card"])
    style.configure("Card.TLabelframe.Label", background=t["bg_card"], foreground=t["fg_title"])
    style.configure("TCombobox", fieldbackground=t["bg_input"], background=t["bg_input"], foreground=t["fg_input"])

    for lbl in form_labels:
        lbl.configure(bg=t["bg_card"], fg=t["fg_label"])

    for entry in (entry_nodes, entry_port_cost, entry_cable_length, entry_cable_cost):
        entry.configure(bg=t["bg_input"], fg=t["fg_input"], insertbackground=t["fg_input"])

    result_label.configure(bg=t["bg_result"], fg=t["fg_result"], insertbackground=t["fg_result"])

    graph_card.configure(bg=t["bg_card"], highlightbackground=t["border_card"])
    graph_card_title.configure(bg=t["bg_card"], fg=t["fg_title"])
    graph_card_sub.configure(bg=t["bg_card"], fg=t["fg_subtitle"])
    graph_frame.configure(bg=t["bg_card"])
    if 'graph_placeholder' in globals() and graph_placeholder.winfo_exists():
        graph_placeholder.configure(bg=t["bg_card"], fg=t["fg_subtitle"])

    if saved_graph is not None:
        show_graph(saved_graph, saved_topology, saved_ring)

def toggle_theme():
    global current_mode
    current_mode = "dark" if current_mode == "light" else "light"
    apply_theme()

# Initialize theme colors
apply_theme()

root.mainloop()
